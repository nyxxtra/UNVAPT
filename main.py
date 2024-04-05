import tkinter as tk
from tkinter import ttk, messagebox
from threading import Thread
from ping import ping_device
from nmap import nmap_scan
from ports.http_80 import scan_port_80  # Import the new function for port 80
from ports.ftp_21 import scan_port_21  # Import the new function for port 21 (FTP)
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# Global variable to track current mode (0 for light, 1 for dark)
current_mode = 0

# Function to perform IP scan
def perform_scan(target_ip):
    # Perform ping and display results
    ping_output = ping_device(target_ip)
    if ping_output:
        messagebox.showinfo("Ping Results", ping_output)

        # Perform Nmap scan and display results
        nmap_output = nmap_scan(target_ip)
        if nmap_output:
            show_custom_window("Nmap Scan Results", nmap_output, target_ip)

# Function to scan IP
def scan_ip():
    target_ip = ip_entry.get()

    # Start the scan in a separate thread
    scan_thread = Thread(target=perform_scan, args=(target_ip,))
    scan_thread.start()

# Function to show custom window with scan results
def show_custom_window(title, message, target_ip):
    # Create a new top-level window
    window = tk.Toplevel()
    window.title(title)
    window.geometry("800x600")  # Set window size

    # Choose color scheme based on current mode
    if current_mode == 0:
        bg_color = "white"
        fg_color = "black"
    else:
        bg_color = "black"
        fg_color = "white"

    # Apply color scheme to the window
    window.config(bg=bg_color)

    # Title label
    title_label = ttk.Label(window, text=title, font=("Times New Roman", 20, "bold"), background=bg_color, foreground=fg_color)
    title_label.pack(pady=10)

    # Create a text widget to display the message
    text_widget = tk.Text(window, wrap=tk.WORD, font=("Times New Roman", 12), background=bg_color, foreground=fg_color)
    text_widget.insert(tk.END, message)
    text_widget.pack(expand=True, fill=tk.BOTH, padx=20, pady=5)

    # Add buttons for exporting to docx and PDF
    export_button_docx = tk.Button(window, text="Export to DOCX", command=lambda: export_to_docx(title, message), font=("Times New Roman", 14), width=15)
    export_button_docx.pack(pady=5)
    export_button_pdf = tk.Button(window, text="Export to PDF", command=lambda: export_to_pdf(title, message), font=("Times New Roman", 14), width=15)
    export_button_pdf.pack(pady=5)

    # Check if port 80 is open and add a button for port 80 scan
    if "80/tcp" in message and "open" in message:
        port_80_button = tk.Button(window, text="Scan Port 80", command=lambda: scan_port_80_window(target_ip), font=("Times New Roman", 14), width=15)
        port_80_button.pack(pady=5)
        
    # Check if port 21 is open and add a button for port 21 scan
    if "21/tcp" in message and "open" in message:
        port_21_button = tk.Button(window, text="Scan Port 21", command=lambda: scan_port_21_window(target_ip), font=("Times New Roman", 14), width=15)
        port_21_button.pack(pady=5)

# Function to open a new window for port 80 scan
def scan_port_80_window(target_ip):
    # Perform a specific scan for port 80 (HTTP)
    port_80_output = scan_port_80(target_ip)
    if port_80_output:
        show_custom_window("Port 80 Scan Results", port_80_output, target_ip)

# Function to open a new window for port 80 scan
def scan_port_21_window(target_ip):
    # Perform a specific scan for port 21 (FTP)
    port_21_output = scan_port_21(target_ip)
    if port_21_output:
        show_custom_window("Port 21 Scan Results", port_21_output, target_ip)
        
# Function to switch between light and dark mode
def switch_mode():
    global current_mode
    current_mode = (current_mode + 1) % 2  # Toggle between light and dark mode

    # Update mode for the main window
    if current_mode == 0:
        # Light mode
        root.config(bg="white")
        ip_label.config(bg="white", fg="black")
        ip_entry.config(bg="white", fg="black")
        scan_button.config(bg="white", fg="black")
        mode_button.config(bg="white", fg="black")
    else:
        # Dark mode
        root.config(bg="black")
        ip_label.config(bg="black", fg="white")
        ip_entry.config(bg="black", fg="white")
        scan_button.config(bg="black", fg="white")
        mode_button.config(bg="black", fg="white")

# Function to export results to DOCX
def export_to_docx(title, message):
    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph(message)
    document.save("scan_results.docx")
    messagebox.showinfo("Export to DOCX", "Scan results exported to scan_results.docx")

# Function to export results to PDF
def export_to_pdf(title, message):
    c = canvas.Canvas("scan_results.pdf", pagesize=letter)
    lines = message.split('\n')
    y = 750
    for line in lines:
        if y < 50:
            c.showPage()
            y = 750
        if "Nmap Scan Results" in line:
            c.setFont("Helvetica", 16)
            c.drawCentredString(300, y, line)
            c.setFont("Helvetica", 12)
        else:
            c.drawString(100, y, line)
        y -= 15
    c.save()
    messagebox.showinfo("Export to PDF", "Scan results exported to scan_results.pdf")

# Create the main window
root = tk.Tk()
root.title("IP Scanner")
root.geometry("400x150")  # Set window size

# Center the window on the screen
def center_window(event=None):
    root.update_idletasks()  # Ensure window dimensions are updated
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()
    x = (screen_width - 400) // 2
    y = (screen_height - 150) // 2
    root.geometry(f"+{x}+{y}")

center_window()  # Center the window initially
root.bind("<Configure>", center_window)  # Center the window when maximized

# Create and place widgets
ip_label = tk.Label(root, text="Enter IP Address:", font=("Times New Roman", 14))
ip_label.grid(row=0, column=0, padx=10, pady=10)

ip_entry = tk.Entry(root, font=("Times New Roman", 14))
ip_entry.grid(row=0, column=1, padx=10, pady=10)

scan_button = tk.Button(root, text="Scan", command=scan_ip, font=("Times New Roman", 14))
scan_button.grid(row=1, column=0, columnspan=2, padx=10, pady=10)

mode_button = tk.Button(root, text="Switch Mode", command=switch_mode, font=("Times New Roman", 14))
mode_button.grid(row=2, column=0, columnspan=2, padx=10, pady=10)

# Run the Tkinter event loop
root.mainloop()
